"""
Generate GM HWIO Deliverable.

This script generates GM HWIO deliverable packages for multiple build targets (APPL, BOOT, RPGM)
based on the GM Supplier Software Technical Requirements document.

The script performs the following main tasks:
    1. Parses compile_commands.json to extract sources, includes, and macros
    2. Builds a dependency graph of source files and their includes
    3. Copies source files to the deliverable directory structure
    4. Processes linker files with appropriate preprocessing
    5. Generates Makefiles from Jinja2 templates
    6. Validates macros and flags against rendered Makefiles
    7. Executes the build script to verify the deliverable compiles

Reference:
    https://magna.sharepoint.com/:w:/r/teams/USF-L2H7890_GM_FPM_VIP_Project/Shared%20Documents/General/KP03%20-%20Product%20Engineering/40_Software/05_Requirements/Supplier_Software_Technical_Requirements%201%202.docx?web=1

TODO:
    - Makefiles should be better generated. Include checks against Macros and Linker Flags.
"""
from __future__ import annotations

import argparse
import filecmp
import importlib.util
import json
import logging
import os
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from codecs import ignore_errors
from collections import defaultdict, deque
from dataclasses import dataclass, field
from functools import lru_cache
from pathlib import Path
from typing import Dict, Iterable, Set, Union

import git
import pypandoc
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell
from filelock import FileLock
from jinja2 import Template

# ---------------------------------------------------------------------------
# Logging / Constants
# ---------------------------------------------------------------------------
logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.WARN)
logging.getLogger("pypandoc").setLevel(logging.ERROR)

IS_WINDOWS = sys.platform.startswith("win")

WIN_GNU_UTILS = Path(r"C:\prjtools\projects\L2H7890_Software\gnu_utils")
WIN_GCC_DIR = Path(r"C:\prjtools\projects\L2H7890_Software\gcc\v15.1.0")
if IS_WINDOWS:
    sys.path += [str(WIN_GNU_UTILS), f"{WIN_GNU_UTILS}/bin", f"{WIN_GCC_DIR}/bin"]

SCRIPT_DIR = Path(__file__).resolve().parent
PREP_DIR = SCRIPT_DIR / "GM_Release_Prep"
DELIVERABLE_DIR = PREP_DIR / "HWIODeliverable"
TEMPLATE_DIR = SCRIPT_DIR / "Templates"

REPO = git.Repo(path=SCRIPT_DIR, search_parent_directories=True)
_project_super = REPO.git.rev_parse("--show-superproject-working-tree").strip()
PROJECT_ROOT = Path(_project_super or REPO.working_tree_dir)

CONFIG_FILE = SCRIPT_DIR / f"{Path(__file__).stem}.config.json"

LFS_ROOT = PROJECT_ROOT / "tools" / "External" / "Blobs"
GNU_UTILS_ZIP = LFS_ROOT / "gnu_utils" / "v1.1.0.zip"

MAP_ANALYSIS_PY = PROJECT_ROOT / "scripts/Map_File_Analysis/Windriver/MapFileAnalysis.py"
_spec = importlib.util.spec_from_file_location("MapFileAnalysis", str(MAP_ANALYSIS_PY))
MapFileAnalysis = importlib.util.module_from_spec(_spec)
sys.modules["MapFileAnalysis"] = MapFileAnalysis
_spec.loader.exec_module(MapFileAnalysis)  # type: ignore

TESTAPP_DIRS: Dict[str, Set[Path]] = {
    "APPL": {
        PROJECT_ROOT / "sw" / "app" / "testapp",
        PROJECT_ROOT / "sw" / "app" / "Features" / "TestApp_Features",
        # PROJECT_ROOT / "sw" / "platform" / "components" / "Test" / "ECC"
    },
    "BOOT": {
        PROJECT_ROOT / "sw" / "boot" / "Test",
    },
    "RPGM": set(),
}
TARGET_SRC_ROOT = {"APPL": "app", "BOOT": "boot", "RPGM": "rpgm"}


# ---------------------------------------------------------------------------
# Small utilities
# ---------------------------------------------------------------------------
def hwio_dir(target: str) -> Path:
    """Return the HWIO deliverable output directory for the given target."""
    return DELIVERABLE_DIR / f"HWIO{target}"


def temp_dir(target: str) -> Path:
    """Return the temporary working directory for the given target."""
    return PREP_DIR / "temp" / f"HWIO{target}"


def ensure_clean_dir(path: Path) -> None:
    """Remove a directory if it exists and recreate it empty."""
    shutil.rmtree(path, ignore_errors=True)
    path.mkdir(parents=True, exist_ok=True)


def create_dir_link(existing_path: Path, new_link_path: Path) -> None:
    """
    Create a directory link (junction on Windows, symlink on Unix).

    Uses Windows junctions which don't require admin privileges.
    Safely replaces any pre-existing file, directory, or symlink at the link path.

    Args:
        existing_path: The target directory that already exists.
        new_link_path: The path where the link will be created.

    Raises:
        NotADirectoryError: If existing_path is not a directory.
        RuntimeError: If mklink command fails on Windows.
    """
    existing_path = existing_path.resolve()
    # Don't resolve() new_link_path to a real location that may not exist yet
    new_link_path = new_link_path

    if not existing_path.is_dir():
        raise NotADirectoryError(f"Target must be a directory: {existing_path}")

    # If something already exists at the link path, remove it robustly
    if new_link_path.exists() or new_link_path.is_symlink():
        try:
            if new_link_path.is_symlink():
                # If it’s already the right target, we’re done
                try:
                    if new_link_path.resolve() == existing_path:
                        return
                except FileNotFoundError:
                    pass
                new_link_path.unlink()
            elif new_link_path.is_dir():
                shutil.rmtree(new_link_path)
            else:
                new_link_path.unlink()
        except FileNotFoundError:
            pass

    if IS_WINDOWS:
        cmd = ["cmd", "/c", "mklink", "/J", str(new_link_path), str(existing_path)]
        result = subprocess.run(cmd, capture_output=True, text=True)
        if result.returncode != 0:
            raise RuntimeError(f"mklink failed: {result.stderr.strip() or result.stdout.strip()}")
    else:
        os.symlink(existing_path, new_link_path, target_is_directory=True)


def copy_file(src: Path, dst: Path, link: bool = False) -> None:
    """
    Copy a file from src to dst, creating parent directories as needed.

    Args:
        src: Source file path.
        dst: Destination file path.
        link: If True, create a hard link instead of copying.

    Note:
        Skips copy if src and dst are the same file or have identical content.
    """
    dst.parent.mkdir(parents=True, exist_ok=True)

    if dst.exists():
        if os.path.samefile(src, dst):
            return

        if filecmp.cmp(src, dst, shallow=True):
            return

    dst.unlink(missing_ok=True)
    os.link(src, dst) if link else shutil.copy2(src, dst)


def read_json(path: Path) -> dict:
    """Read and parse a JSON file, returning its contents as a dictionary."""
    return json.loads(path.read_text())


def write_json(path: Path, data: dict) -> None:
    """Write a dictionary to a JSON file with indentation, creating parent dirs."""
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(data, indent=4))


def norm_sorted(items: Iterable[str]) -> list[str]:
    """Return a case-insensitive sorted list of unique strings."""
    return sorted(set(items), key=lambda s: s.lower())


# ---------------------------------------------------------------------------
# compile_commands.json parsing
# ---------------------------------------------------------------------------
def parse_compile_commands_json(json_file: Path, out_json: Path) -> dict[str, list[str]]:
    """
    Parse a compile_commands.json and extract source files, includes, and macros.

    Reads the CMake-generated compile_commands.json, extracts relevant compiler
    flags, and categorizes them into sources, includes, macros (regular, used,
    partition, test), and other flags.

    Args:
        json_file: Path to the compile_commands.json file.
        out_json: Path where the parsed output JSON will be written.

    Returns:
        Dictionary with keys: sources, includes, macros, macros_used,
        macros_partition, macros_test, other_flags.
    """
    data = json.loads(json_file.read_text())

    srcs, incs = set(), set()
    macros, macros_used, macros_part, macros_test, other = set(), set(), set(), set(), set()

    for item in data:
        src_rel = Path(item["file"]).as_posix().split("/sw/", 1)[-1]
        srcs.add(src_rel)

        cmd = item["command"]
        filtered = {
            tok
            for tok in cmd.split(" ")
            if not any(
                [
                    tok in {"-c", "-o"},
                    tok.startswith(("/", "\\")),
                    tok.endswith((".obj", ".o", ".c", ".s", ".exe")),
                    "CMAKE" in tok,
                    not tok,
                ]
            )
        }

        for tok in filtered:
            if tok.startswith("-I"):
                inc = Path(tok).as_posix().removeprefix("-I").split("/sw/", 1)[-1]
                incs.add(inc)
            elif tok.startswith("-D"):
                m = tok.removeprefix("-D")
                if m.endswith("_USED"):
                    macros_used.add(m)
                elif m.endswith("_Config_EcucPartition"):
                    macros_part.add(m)
                elif m.startswith("TEST_"):
                    macros_test.add(m)
                else:
                    macros.add(m)
            else:
                other.add(tok)

    payload = {
        "sources": norm_sorted(srcs),
        "includes": norm_sorted(incs),
        "macros": norm_sorted(macros),
        "macros_used": norm_sorted(macros_used),
        "macros_partition": norm_sorted(macros_part),
        "macros_test": norm_sorted(macros_test),
        "other_flags": norm_sorted(other),
    }
    write_json(out_json, payload)
    return payload


# ---------------------------------------------------------------------------
# Source graph
# ---------------------------------------------------------------------------
@dataclass(eq=False)
class SourceFile:
    """
    Represents a source file in the dependency graph.

    Tracks the file path, which include keys were matched during resolution,
    and child SourceFile nodes representing included headers. Used to build
    a transitive closure of all dependencies for a compilation unit.

    Attributes:
        path: Absolute path to the source or header file.
        include_keys_seen: Set of include strings that resolved to this file.
        children: Set of SourceFile nodes this file includes.
    """

    path: Path
    include_keys_seen: Set[str] = field(default_factory=set)
    children: Set["SourceFile"] = field(default_factory=set, compare=False)

    @property
    @lru_cache(maxsize=None)
    def include_keys(self) -> Set[str]:
        """Extract all #include directives from this file's source text."""
        text = self.path.read_text(errors="ignore")
        pat = re.compile(r'#\s*include\s+["<]([^">]+)[">]')
        return {m.group(1) for m in pat.finditer(text)}

    @property
    def include_dirs_for_this(self) -> Set[Path]:
        """Derive include directories from resolved include keys."""
        if not self.include_keys_seen:
            return set()
        path_str = self.path.as_posix()
        return {Path(path_str.removesuffix(k)) for k in self.include_keys_seen}

    def all_deps(self, visited: Set[Path] | None = None) -> Set[Path]:
        """Return all transitive file dependencies (paths) via BFS."""
        visited = visited or set()
        if self.path in visited:
            return visited
        visited.add(self.path)
        for c in self.children:
            visited |= c.all_deps(visited)
        return visited

    def all_include_dirs(self, visited: Set[Path] | None = None) -> Set[Path]:
        """Return all transitive include directories from this file and its dependencies."""
        visited = visited or set()
        if self.path in visited:
            return self.include_dirs_for_this
        visited.add(self.path)
        incs = set(self.include_dirs_for_this)
        for c in self.children:
            incs |= c.all_include_dirs(visited)
        return incs

    def __hash__(self) -> int:
        return hash(self.path.resolve())

    def __eq__(self, other: object) -> bool:
        return isinstance(other, SourceFile) and self.path.resolve() == other.path.resolve()


# ---------------------------------------------------------------------------
# File collection
# ---------------------------------------------------------------------------
def retrieve_wanted_source_files(source_root: Path, dst_root: Path, parsed_json_path: Path) -> dict:
    """
    Collect source files and their dependencies, then copy them to the deliverable.

    Builds a dependency graph from the parsed compile_commands.json, resolves
    transitive includes, separates private library sources, and copies all
    needed files to the appropriate destination directories.

    Args:
        source_root: Root path of the sw/ source tree.
        dst_root: Destination root for copied Source files.
        parsed_json_path: Path to the parsed compile_commands JSON.

    Returns:
        Dictionary with 'sources' (main source files) and 'libs' (private lib files).
    """
    cfg = read_json(CONFIG_FILE)
    test_dirs = cfg.get("test_dirs")
    private_libs = cfg.get("private_libs")

    parsed = read_json(parsed_json_path)
    src_rel_list: list = parsed["sources"]
    include_rel_list = set(parsed["includes"])
    source_root = source_root.resolve(strict=True)

    # RPGM build exceptions
    if target == "RPGM":
        drop_files = {
            'Can_17_McmCan_Irq.c',
            'Ifx_Ssw_Tc1.c',
            'Ifx_Ssw_Tc2.c',
            'Ifx_Ssw_Tc3.c',
            'Ifx_Ssw_Tc4.c',
            'Ifx_Ssw_Tc5.c',
            'Ifx_Ssw_Tc6.c',
            'SchM_Can_17_McmCan.c',
            'SchM_Dma.c',
            'SchM_Gpt.c',
            'SchM_Icu.c',
            'SchM_Pmic.c',
            'SchM_Spi.c',
            'SchM_Spi.c',
        }
        src_rel_list = [
            s for s in src_rel_list if not any(df.lower() in s.lower() for df in drop_files)
        ]

    # Seed wanted sources
    wanted: Set[SourceFile] = {SourceFile(path=source_root / s) for s in src_rel_list}
    test_folders = TESTAPP_DIRS[target]
    for test_folder in test_folders:
        wanted |= {SourceFile(path=p.resolve()) for p in test_folder.rglob("*.c")}
        wanted |= {SourceFile(path=p.resolve()) for p in test_folder.rglob("*.s")}

    # Include roots = declared + parent dirs of seed sources
    include_paths: Set[Path] = {source_root / p for p in include_rel_list} | {s.path.parent for s in wanted}
    include_paths = _validate_include_roots(source_root, include_paths)

    # Index headers under include roots
    rel_index: Dict[str, list[Path]] = defaultdict(list)
    for root in list(include_paths):
        if not root.exists():
            continue
        for h in root.rglob("*.h"):
            rel_index[h.relative_to(root).as_posix()].append(h)

    # Resolve transitive includes
    multiple_hits, no_hits = set(), set()
    seen: Set[SourceFile] = set()
    file_map: Dict[Path, SourceFile] = {s.path: s for s in wanted}
    q = deque(wanted)

    while q:
        cur = q.popleft()
        if cur in seen:
            continue
        seen.add(cur)
        for key in cur.include_keys:
            hits = [p for p in rel_index.get(key, []) if any(p.resolve() == (root / key).resolve() for root in include_paths)]
            if not hits:
                no_hits.add(key)
                continue
            if len(hits) > 1:
                multiple_hits.add(key)
            for hp in hits:
                if hp not in file_map:
                    file_map[hp] = SourceFile(path=hp)
                node = file_map[hp]
                node.include_keys_seen.add(key)
                cur.children.add(node)
                if node not in seen:
                    q.append(node)

    logger.warning(f"No include files found ({len(no_hits)}) {no_hits}")
    logger.warning(f"Multiple include files found ({len(multiple_hits)}) {multiple_hits}")

    # Filter to .c/.s
    only_cs = {p: s for p, s in file_map.items() if p.suffix in {".c", ".s"}}

    # Split out private libs
    libs: Dict[str, dict] = {}
    for lib_name, lib_dirs in private_libs.items():
        lib_paths = tuple(str((source_root / d).resolve()) for d in lib_dirs)
        libs[lib_name] = {
            "files": {
                p: s
                for p, s in only_cs.items()
                if str(p).startswith(lib_paths) and not ("testapp" in p.as_posix().lower())
            }
        }
        only_cs = {p: s for p, s in only_cs.items() if p not in libs[lib_name]["files"]}  # Remove from main

    # Copy remaining + lib deps
    _copy_sources_with_tests(only_cs, source_root, dst_root, test_dirs)
    for lib_name, lib in libs.items():
        logger.info(f"Copying private lib: {lib_name}")
        for s in lib["files"].values():
            for dep in s.all_deps():
                dst = dst_root.parent.parent.parent / "Private" / "Source" / dep.relative_to(source_root)
                copy_file(dep, dst)

    return {"sources": only_cs, "libs": libs}


def _validate_include_roots(source_root: Path, include_roots: Set[Path]) -> Set[Path]:
    """
    Validate include root directories exist and have correct capitalization.

    Args:
        source_root: The root of the source tree to search.
        include_roots: Set of include directory paths to validate.

    Returns:
        Set of validated include root paths that exist.

    Raises:
        ValueError: If any include paths have incorrect capitalization.
    """
    ok, missing, wrong_caps = set(), set(), False
    all_dirs = {p for p in source_root.rglob("*") if p.is_dir()} | {source_root}
    lookup = {p.resolve().as_posix().lower(): p for p in all_dirs}

    for inc in include_roots:
        k = inc.resolve().as_posix().lower()
        if k in lookup:
            if lookup[k].as_posix() != inc.resolve().as_posix():
                logger.error(
                    "Include path exists with different capitalization\n\t"
                    f"Expected: {inc.as_posix()}\n\tFound: {lookup[k].as_posix()}"
                )
                wrong_caps = True
            ok.add(lookup[k])
        else:
            if not any(["asr" in str(inc), str(source_root) not in str(inc)]):
                missing.add(inc)

    if missing:
        logger.warning("Missing include paths:\n\t- " + "\n\t- ".join(str(p) for p in sorted(missing)))
    if wrong_caps:
        raise ValueError("Include paths have wrong capitalization.")
    return ok


def _copy_sources_with_tests(src_map: Dict[Path, SourceFile], root: Path, dst_root: Path, test_dirs: list[str]) -> None:
    """
    Copy source files and dependencies, routing test files to a Test directory.

    Args:
        src_map: Dictionary mapping paths to SourceFile objects.
        root: Source tree root path.
        dst_root: Destination root for Source files.
        test_dirs: List of relative paths that identify test directories.
    """
    deps = {f for s in src_map.values() for f in s.all_deps()}
    for dep in deps:
        is_test = any(str(dep).startswith(str((root / td).resolve())) for td in test_dirs)
        dst = (dst_root.parent / "Test" / dep.relative_to(root)) if is_test else (dst_root / dep.relative_to(root))
        copy_file(dep, dst)


# ---------------------------------------------------------------------------
# Linker handling
# ---------------------------------------------------------------------------
def handle_linker_files() -> None:
    """
    Process and split linker script files for the current target.

    Preprocesses the main linker script using GCC, extracts marked sections
    into separate include files (symbols, memory ranges, sections), and
    cleans up the output. Creates both main linker file and supplier-specific
    test/CAL section files.
    """
    def capture_linker_tree(entry: Path, out_dir: Path) -> None:
        """Recursively copy linker includes, adding section markers to each file."""
        def add_markers(p: Path) -> None:
            """Wrap file content with LINKER_FLAG START/END markers."""
            txt = p.read_text()
            p.write_text(f"/* LINKER_FLAG START: {p.name} */\n{txt}\n/* LINKER_FLAG END: {p.name} */")

        q = deque([entry])
        pat = re.compile(r'#\s*include\s+["<]([^">]+)[">]')
        seen = set()
        while q:
            cur = q.popleft()
            if cur in seen:
                continue
            seen.add(cur)
            txt = cur.read_text()
            for m in pat.finditer(txt):
                inc_name = m.group(1)
                inc_path = cur.parent / inc_name
                q.append(inc_path)
                txt = txt.replace(m.group(0), f'#include "{inc_path.name}"')
            out = out_dir / cur.name
            out.write_text(txt)
            add_markers(out)

    def extract_or_replace(parse_file: Union[str, Path], name: str, replace: Union[str, bool] = False) -> str:
        """
        Extract text between LINKER_FLAG markers, optionally replacing it.

        Args:
            parse_file: Path to the file to parse.
            name: The marker name to search for (e.g., 'symbols_appl.lcf').
            replace: If a string, replace the marked section with this text.

        Returns:
            The text content between the START and END markers.
        """
        p = Path(parse_file)
        text = p.read_text()
        s_pat = re.compile(rf"/\*\s*LINKER_FLAG\s+START:\s*{re.escape(name)}\s*\*/")
        e_pat = re.compile(rf"/\*\s*LINKER_FLAG\s+END:\s*{re.escape(name)}\s*\*/")
        ms, me = s_pat.search(text), e_pat.search(text)
        if not ms or not me:
            logger.info(f"Markers for {name} not found in {p} for target {target}")
            return ""
        inner = text[ms.end(): me.start()]
        if replace is not False:
            p.write_text(text[: ms.start()] + replace + text[me.end():])
        return inner

    logger.info(f"Handling the linker files for target: {target}")

    tdir = temp_dir(target)
    out_hwio = hwio_dir(target)

    linker_src_dir = PROJECT_ROOT / "sw" / "common" / "linker"
    linker_tmp_dir = tdir / "__linker"
    sections_test = linker_src_dir / f"sections_test_{target.lower()}.lcf"
    sections_cal = linker_src_dir / "sections_CAL.lcf"

    tmp_main = tdir / f"preprocessed_{target}.lsl"
    out_main = out_hwio / f"HWIO{target}.lsl"

    path_linker_includes = out_hwio / "LinkerIncludes"
    path_linker_includes.mkdir(exist_ok=True, parents=True)

    out_symbols = path_linker_includes / f"symbols_{target.lower()}.lcf"
    out_mem = path_linker_includes / "mem_ranges.lcf"
    out_sections = path_linker_includes / f"sections_{target.lower()}.lcf"
    out_sections_test = DELIVERABLE_DIR / "Supplier" / f"HWIO{target}" / sections_test.name
    out_sections_cal = DELIVERABLE_DIR / "Supplier" / f"HWIO{target}" / sections_cal.name

    # Prep
    out_main.parent.mkdir(parents=True, exist_ok=True)
    tmp_main.parent.mkdir(parents=True, exist_ok=True)
    out_sections_test.parent.mkdir(parents=True, exist_ok=True)
    ensure_clean_dir(linker_tmp_dir)

    # Pre-parse tree
    main_in = linker_src_dir / f"Lcf_Windriver_Tricore_Tc_4D7_{target}.lsl"
    capture_linker_tree(main_in, linker_tmp_dir)

    gcc = "gcc" if not IS_WINDOWS else str(WIN_GCC_DIR / "bin" / "gcc.exe")
    subprocess.run(
        [gcc, "-E", "-P", "-C", "-x", "c", "-nostdinc", "-undef", str(linker_tmp_dir / main_in.name), "-o", str(tmp_main)],
        check=True,
    )

    text = extract_or_replace(tmp_main, sections_test.name, replace=f'\n/* #include "./../Supplier/HWIO{target}/sections_test_{target.lower()}.lcf" */\n')
    out_sections_test.write_text(text)

    if target == "APPL":
        text = extract_or_replace(tmp_main, sections_cal.name, replace=f'\n  #include "./../Supplier/HWIO{target}/sections_CAL.lcf"\n')
        out_sections_cal.write_text(f"#ifndef LINK_GM_LIB\n{text}\n#endif\n")

    text = extract_or_replace(tmp_main, out_symbols.name, replace=f'#include "LinkerIncludes/{out_symbols.name}"')
    if target == "APPL":
        text = re.sub(r".*HWIOAPPL_fpu_hwio\.a.*\n", "", text)
    out_symbols.write_text(text)

    text = extract_or_replace(tmp_main, out_mem.name, replace=f'  #include "LinkerIncludes/{out_mem.name}"')
    out_mem.write_text(text)

    text = extract_or_replace(tmp_main, out_sections.name, replace=f'  #include "LinkerIncludes/{out_sections.name}"')
    out_sections.write_text(text)

    # Strip block comments & collapse whitespace
    pat_block = re.compile(r"/\*.*?\*/\n*", re.DOTALL)
    cleaned = pat_block.sub("\n", tmp_main.read_text())
    cleaned = re.sub(r"(\s*\n)+", "\n", cleaned).strip() + "\n"

    # Comment out the supplier items
    cleaned = re.sub(r'(#include.*?Supplier.*")', r"/* \1 */", cleaned, flags=re.IGNORECASE)

    tmp_main.write_text(cleaned)
    tmp_main.rename(out_main)


# ---------------------------------------------------------------------------
# Static file moves
# ---------------------------------------------------------------------------
def copy_static_files() -> None:
    """
    Copy static files and libraries required for the deliverable.

    Handles RTAOS library for APPL, customer libraries, iLLD source relocation,
    cleanup of unwanted file types, and common_env.ps1 modification for gnu_utils.
    Also triggers linker file processing.
    """
    out_hwio = hwio_dir(target)
    rtaos_rel = "app/asr/Os/Implementation/RTAOS.a"
    rtaos_dst = out_hwio / "Source" / rtaos_rel
    cust_lib = PROJECT_ROOT / "sw" / "customer" / target / f"GM_{target}.lib"
    cust_lib_dst = DELIVERABLE_DIR / "Customer" / target / cust_lib.name
    common_env_ps1 = PROJECT_ROOT / "scripts" / "common_env.ps1"

    if target == "APPL":
        copy_file(PROJECT_ROOT / "sw" / rtaos_rel, rtaos_dst)

    if cust_lib.exists():
        copy_file(cust_lib, cust_lib_dst)
        wanted = cust_lib.parent / "wanted.list"
        copy_file(wanted, cust_lib_dst.parent / wanted.name)

    # iLLD exception → move C to Test, copy headers
    illd_src = out_hwio / "Source" / "platform" / "iLLD"
    for f in illd_src.rglob("*.*"):
        if not f.is_file() or f.suffix not in {".c", ".h", ".s"}:
            continue
        dst = Path(f.as_posix().replace("Source", "Test"))
        dst.parent.mkdir(parents=True, exist_ok=True)
        (f.rename(dst) if f.suffix == ".c" else copy_file(f, dst))

    # Cleanup junk + TC46xA
    for f in set((out_hwio / "Test").rglob("*.*")) | set((out_hwio / "Source").rglob("*.*")):
        if f.suffix not in {".c", ".h", ".s", ".a"} or "TC46xA" in f.parts:
            f.unlink(missing_ok=True)

    # Drop-in common_env with gnu_utils PATH
    env_txt = common_env_ps1.read_text()
    env_txt += '\nAdd-ToPath -NewPath "$PSScriptRoot/gnu_utils/v1.1.0"'
    env_txt += '\nAdd-ToPath -NewPath "$PSScriptRoot/gnu_utils/v1.1.0/bin"'
    (PREP_DIR / "common_env.ps1").write_text(env_txt)

    handle_linker_files()


# ---------------------------------------------------------------------------
# Templates / docs
# ---------------------------------------------------------------------------
def _mkvar(name: str, items: Iterable[str], prefix: str = "") -> str:
    """
    Format a Makefile variable assignment with line continuations.

    Args:
        name: The variable name (e.g., 'HWIO_SRCS').
        items: Iterable of values to assign.
        prefix: Optional prefix for each value (e.g., '-I' for includes).

    Returns:
        Formatted Makefile variable string with proper line continuations.
    """
    backslash = '\\'
    vals = norm_sorted(items)
    if not vals:
        return f"{name} ="
    body = [f"{name} = \\"]
    for i, v in enumerate(vals):
        body.append(f"\t{prefix}{v}{' ' + backslash if i < len(vals) - 1 else ''}")
    return "\n".join(body)


def render_jinja_templates(source_tree: dict) -> None:
    """
    Generate Makefiles from Jinja2 templates based on collected source files.

    Creates target-specific Makefiles with source files, include directories,
    and test files. Also generates separate Makefiles for private libraries.

    Args:
        source_tree: Dictionary containing 'sources' and 'libs' mappings.
    """
    cfg = read_json(CONFIG_FILE)
    test_dirs = cfg.get("test_dirs")
    sw_root = (PROJECT_ROOT / "sw").resolve()
    tmpl = (TEMPLATE_DIR / "template.target.Makefile").read_text()

    def is_test(p: Path) -> bool:
        """Check if a path is under any of the configured test directories."""
        rp = p.resolve()
        return any(str(rp).startswith(str((sw_root / td).resolve())) for td in test_dirs)

    def rel_source_or_test(p: Path) -> str:
        """Return relative path prefixed with ./Source or ./Test as appropriate."""
        tag = "Test" if is_test(p) else "Source"
        return f"./{tag}/{p.resolve().relative_to(sw_root).as_posix()}"

    def rel_private(p: Path) -> str:
        """Return relative path for private library files (./Source/<sw-relative>)."""
        # libs were copied to ./Private/Source/<sw-relative>
        return f"./Source/{p.resolve().relative_to(sw_root).as_posix()}"

    def collect_includes(files_map: Dict[Path, SourceFile]) -> Set[Path]:
        """Collect all include directories from a set of SourceFile objects."""
        inc = set()
        for s in files_map.values():
            inc |= s.all_include_dirs()
        return inc

    def sanitize_name(name: str) -> str:
        """Convert a name to uppercase with non-alphanumeric chars replaced by underscores."""
        return re.sub(r"[^A-Za-z0-9]", "_", name).upper()

    # ---------- MAIN (split Source/Test) ----------
    main_map: Dict[Path, SourceFile] = source_tree["sources"]
    main_test_map = {p: s for p, s in main_map.items() if is_test(p)}
    main_src_map = {p: s for p, s in main_map.items() if p not in main_test_map}

    main_src_files = {rel_source_or_test(p) for p in main_src_map}
    main_test_files = {rel_source_or_test(p) for p in main_test_map}

    main_inc_dirs_abs = {p for p in collect_includes(main_src_map) if not is_test(p)}
    main_test_inc_abs = {p for p in collect_includes(main_test_map) if p not in main_inc_dirs_abs}

    main_inc_dirs = {rel_source_or_test(p) for p in main_inc_dirs_abs}
    main_test_inc = {rel_source_or_test(p) for p in main_test_inc_abs}
    main_test_inc.add("./Test/platform/components/Test/ECC")  # legacy exception

    ctx = {
        "lib_only": False,
        "libs": "",
        'lib_name': f"HWIO{target}",
        "TARGET": f"HWIO{target}",
        "TARGET_lower": f"hwio{target.lower()}",
        "TARGET_short": target,
        "TARGET_short_lower": target.lower(),
        "HWIO_SRCS": _mkvar("HWIO_SRCS", main_src_files),
        "INCLUDES": _mkvar("INCLUDES", main_inc_dirs, prefix="-I"),
        "TEST_SRCS": _mkvar("TEST_SRCS", main_test_files),
        "TEST_INCLUDES": _mkvar("TEST_INCLUDES", main_test_inc, prefix="-I"),
    }

    # ---------- LIBS (NO split; everything under Private/Source) ----------
    for lib_name, lib_info in source_tree.get("libs", {}).items():
        files_map: Dict[Path, SourceFile] = lib_info.get("files", {})
        if not files_map:
            continue
        ctx["libs"] += f" ./Products/HWIO{target}_{sanitize_name(lib_name)}.a"

        lib_src_files = {rel_private(p) for p in files_map.keys()}
        lib_inc_dirs_abs = collect_includes(files_map)
        lib_inc_dirs = {rel_private(p) for p in lib_inc_dirs_abs}

        new_items = {
            "lib_only": True,
            "libs": "",
            'lib_name': f'HWIO{target}_{sanitize_name(lib_name)}',
            "TARGET": f"HWIO{target}",
            "TARGET_lower": f"hwio{target.lower()}",
            "TARGET_short": target,
            "TARGET_short_lower": target.lower(),
            f"HWIO_SRCS": _mkvar(f"HWIO_SRCS", lib_src_files),
            f"INCLUDES": _mkvar(f"INCLUDES", lib_inc_dirs, prefix="-I"),
            f"TEST_SRCS": "",
            f"TEST_INCLUDES": "",
        }

        lib_makefile_dest = PREP_DIR / "Private" / f"{target}_{lib_name}.makefile"
        lib_makefile_dest.parent.mkdir(parents=True, exist_ok=True)
        lib_makefile_dest.write_text(Template(tmpl).render(**new_items))

    out_dir = DELIVERABLE_DIR / f"HWIO{target.upper()}"
    out_dir.mkdir(parents=True, exist_ok=True)
    (out_dir / "Makefile").write_text(Template(tmpl).render(**ctx))


def render_static_templates() -> None:
    """
    Copy static template files to the deliverable directory.

    Copies the top-level Makefile, common.mk, and build script. Also converts
    documentation from Markdown to DOCX and plain text formats using pypandoc.
    """
    # top-level Makefile
    top_mk = DELIVERABLE_DIR / "Makefile"
    top_mk.unlink(missing_ok=True)
    copy_file(TEMPLATE_DIR / "template.Makefile", top_mk)

    common_mk = DELIVERABLE_DIR / "common.mk"
    common_mk.unlink(missing_ok=True)
    copy_file(TEMPLATE_DIR / "template.common.mk", common_mk)

    # build script (only if not already a custom one)
    build_ps1 = PREP_DIR / "build.ps1"
    templ_build = TEMPLATE_DIR / "template.build.ps1"
    if build_ps1.exists() and build_ps1.resolve() != templ_build.resolve():
        build_ps1.unlink(missing_ok=True)
    if not build_ps1.exists():
        copy_file(templ_build, build_ps1)

    # Patch for GM Lib build
    patch_dest = PREP_DIR / "Private" / "APPL_GM_BUILD.patch"
    patch_template = TEMPLATE_DIR / patch_dest.name
    copy_file(src=patch_template, dst=patch_dest)

    # docs
    pypandoc.convert_file(str(TEMPLATE_DIR / "progman.md"), to="docx", outputfile=str(DELIVERABLE_DIR / "progman.docx"))
    pypandoc.convert_file(str(TEMPLATE_DIR / "release.md"), to="plain", outputfile=str(DELIVERABLE_DIR / "release.txt"))


def capture_docs() -> None:
    """
    Convert and copy documentation files to the deliverable.

    Converts the SPDS reference Markdown document to DOCX format and applies
    formatting fixes via reformat_docx.
    """
    docs = DELIVERABLE_DIR / "docs"
    docs.mkdir(parents=True, exist_ok=True)
    spds_md = PROJECT_ROOT / "sw/app/hwio/spds_hwio/docs/Spds_Reference_Doc.md"
    pypandoc.convert_file(str(spds_md), to="docx", outputfile=str(docs / "Spds_Reference_Doc.docx"))
    reformat_docx(docs / "Spds_Reference_Doc.docx")

    flash_integration_notes = PROJECT_ROOT / "docs/External/FPM2_VIP_External_Flash_IntegrationNotes.pdf"
    copy_file(flash_integration_notes, docs / flash_integration_notes.name)


def reformat_docx(docx_path: Path) -> None:
    """
    Apply formatting fixes to a Word document.

    Adds borders to all table cells and highlights any TODO text in yellow.

    Args:
        docx_path: Path to the DOCX file to reformat.
    """
    def set_cell_border(cell: _Cell, **kwargs) -> None:
        """Apply border styling to a table cell using OOXML elements."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement("w:tcBorders")
            tcPr.append(tcBorders)
        for edge in ("start", "top", "end", "bottom", "insideH", "insideV"):
            data = kwargs.get(edge)
            if not data:
                continue
            tag = f"w:{edge}"
            el = tcBorders.find(qn(tag)) or OxmlElement(tag)
            if el not in tcBorders:
                tcBorders.append(el)
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in data:
                    el.set(qn(f"w:{key}"), str(data[key]))

    doc = Document(str(docx_path))
    for table in doc.tables:
        table.autofit = True
        for row in table.rows:
            for cell in row.cells:
                set_cell_border(
                    cell,
                    top={"sz": 4, "val": "single", "color": "#000000", "space": "0"},
                    bottom={"sz": 4, "color": "#000000", "val": "single"},
                    start={"sz": 4, "val": "single", "color": "#000000", "space": "0"},
                    end={"sz": 4, "val": "single", "color": "#000000", "space": "0"},
                )
    for para in doc.paragraphs:
        for run in para.runs:
            if "TODO" in run.text:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    doc.save(str(docx_path))


# ---------------------------------------------------------------------------
# Consistency checks / build
# ---------------------------------------------------------------------------
def check_macros_in_source_and_rendered(src_json: Path, makefiles: Set[Path]) -> None:
    """
    Validate that all required macros and flags appear in the generated Makefiles.

    Compares the macros extracted from compile_commands.json against the
    rendered Makefile content and raises an error if any required flags are missing.

    Args:
        src_json: Path to the parsed compile_commands JSON.
        makefiles: Set of Makefile paths to check against.

    Raises:
        ValueError: If extra sections exist in JSON or required flags are missing.
    """
    parsed = read_json(src_json)
    ignore_sections = {"sources", "includes", "macros_used"}
    wanted = {"macros", "macros_partition", "macros_test", "other_flags"}

    extras = parsed.keys() - ignore_sections - wanted
    if extras:
        raise ValueError(f"Extra sections in parsed_compile_commands.json: {extras}")

    allowed_false_neg = {
        "APP_SW=DCU_SW",
        "APP_SW=APPL",
        "APP_SW=BOOT",
        "APP_SW=RPGM",
        "APPL_BUILD_TYPE=Debug",
        "BOOT_BUILD_TYPE=Debug",
        "RPGM_BUILD_TYPE=Debug",
        "REDECLARED"
    }
    wanted_flags = set().union(*(parsed[k] for k in wanted))
    mk_text = "\n".join(txt.read_text() for txt in makefiles)

    missing = {f for f in wanted_flags if f not in mk_text} - allowed_false_neg
    if missing:  # TODO: This is failing..? but only on Jenkins?
        raise ValueError(f"Missing flags in Makefile: {missing} for target {target}")


def run_build_script(build_script: Path, build_target: str) -> None:
    """
    Execute the PowerShell build script for the specified target.

    Args:
        build_script: Path to the build.ps1 script.
        build_target: The build target name (APPL, BOOT, or RPGM).

    Raises:
        RuntimeError: If the build script exits with a non-zero return code.
    """
    logger.info(f"Running the build script {build_script}")
    res = subprocess.run(["pwsh", "-NoLogo", "-NoProfile", "-NonInteractive", "-File", str(build_script), build_target],
                         stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    if res.returncode != 0:
        logger.error(res.stdout.decode())
        logger.error(res.stderr.decode())
        raise RuntimeError(f"Build script `{build_script}` failed with return code {res.returncode} for target {build_target}")


def pull_required_bins() -> None:
    """
    Pull and extract GNU utilities required for the build.

    Uses git-lfs to pull the gnu_utils archive, extracts it, and creates
    a directory link in the prep directory. Uses a file lock to prevent
    concurrent extraction attempts.
    """
    tools = PROJECT_ROOT / "tools"
    tools_repo = git.Repo(path=tools, search_parent_directories=True)

    gnu_dir = Path(str(GNU_UTILS_ZIP).removesuffix(".zip"))
    with FileLock(lock_file="gnu_utils_lfs.lock", timeout=30):
        if not gnu_dir.exists():
            try:
                tools_repo.git.lfs("pull", "--include", str(GNU_UTILS_ZIP.relative_to(tools)))
            except git.GitCommandError as e:
                logger.error(f"Failed to pull gnu_utils via git lfs: {e}")
            with zipfile.ZipFile(GNU_UTILS_ZIP, "r") as zf:
                zf.extractall(gnu_dir)

    target_link = PREP_DIR / "gnu_utils"
    if (not target_link.exists()) or (target_link.resolve() != gnu_dir.resolve()):
        target_link.unlink(missing_ok=True)
        create_dir_link(gnu_dir, target_link)


# ---------------------------------------------------------------------------
# Orchestration
# ---------------------------------------------------------------------------
def build_one_target() -> None:
    """
    Build the deliverable for a single target (APPL, BOOT, or RPGM).

    Orchestrates the full build process: parses compile commands, collects
    source files, copies static assets, generates Makefiles, and validates
    macros. Uses the global ``target`` variable set in the main loop.
    """
    build_dir = PROJECT_ROOT / "build" / target
    out_hwio = hwio_dir(target)
    tdir = temp_dir(target)

    DELIVERABLE_DIR.mkdir(parents=True, exist_ok=True)
    tdir.mkdir(parents=True, exist_ok=True)

    ensure_clean_dir(out_hwio)  # clean per-target output container parent already exists

    parsed_path = tdir / "parsed_compile_commands.json"
    parse_compile_commands_json(build_dir / "compile_commands.json", parsed_path)

    src_tree = retrieve_wanted_source_files(PROJECT_ROOT / "sw", out_hwio / "Source", parsed_path)

    copy_static_files()

    # ECC exception cleanup (unchanged logic)
    ecc_src = out_hwio / "Source" / "platform" / "components" / "Test" / "ECC"
    for c in ecc_src.rglob("*.c"):
        c.unlink(missing_ok=True)

    render_jinja_templates(src_tree)
    check_macros_in_source_and_rendered(
        src_json=parsed_path,
        makefiles={
            out_hwio / "Makefile",
            DELIVERABLE_DIR / "common.mk",
        }
    )


def delete_empty_folders_recursively(root: Path) -> None:
    """
    Remove all empty directories under root, bottom-up.

    Skips directories named 'tmp'. Uses os.walk in reverse order to ensure
    child directories are processed before parents.

    Args:
        root: The root directory to clean.
    """
    for r, dirs, files in os.walk(root, topdown=False):
        for d in dirs:
            full = Path(r) / d
            if not any(full.iterdir()):
                if d in {"tmp"}:
                    continue
                os.rmdir(full)


def replace_copyright_headers(root: Path) -> None:
    """
    Update copyright headers in source files for GM deliverable.

    Replaces the generic Magna copyright text with GM-specific wording
    in all .c, .h, and .s files under the root directory.

    Args:
        root: Root directory containing source files to update.
    """
    exts = {".c", ".h", ".s"}
    needle = "the purpose of the specific program or project in which"
    repl = "the purposes of the General Motors VIP FPM program in which"
    for p in (p for p in root.rglob("*") if p.suffix in exts):
        haystack = p.read_text(encoding="utf-8", errors="ignore")
        if needle in haystack:
            with tempfile.NamedTemporaryFile("w", dir=p.parent, delete=False, errors="ignore", encoding="utf-8") as tmp:
                tmp.write(haystack.replace(needle, repl))
            os.replace(tmp.name, p)  # Atomic writing to minimize race conditions.


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    targets_all = {"APPL", "BOOT", "RPGM"}

    ap = argparse.ArgumentParser(description="Generate GM Deliverable")
    ap.add_argument("--target", action="store", type=str.upper, default="ALL", choices=targets_all | {"ALL"}, help="Target to build.")
    args = ap.parse_args()

    DELIVERABLE_DIR.mkdir(parents=True, exist_ok=True)
    render_static_templates()
    capture_docs()
    pull_required_bins()

    to_process = targets_all if args.target == "ALL" else {args.target}
    for target in to_process:  # NOTE: global name used by helpers (kept by design)
        logger.info(f"Processing target: {target}")
        build_one_target()
        run_build_script(build_script=PREP_DIR / "build.ps1", build_target=target)
        # MapFileAnalysis.main(
        #     file_map=hwio_dir(target) / 'Products' / f'Magna_Test{target.lower()}.map',
        #     output_folder=hwio_dir(target) / 'Products'
        # )
        delete_empty_folders_recursively(hwio_dir(target))
    replace_copyright_headers(DELIVERABLE_DIR)
